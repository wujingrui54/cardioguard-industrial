# -*- coding: utf-8 -*-
"""
convert_paper.py
================
把 PAPER_ZH.md / PAPER_KR.md / PAPER_EN.md 转成 Word .docx,
便于竞赛官方系统/审稿人离线阅读。

依赖:
    pip install python-docx markdown

注:
- 三种语言都使用对应区域字体: 中文 = 等线 / 微软雅黑, 韩文 = Malgun Gothic, 英文 = Calibri / Times New Roman
- 自动建表头/表单
- 标题层级 # / ## / ### 自动映射 Heading 1/2/3

运行:
    python convert_paper.py

产出:
    paper/CardioGuard_paper_ZH.docx
    paper/CardioGuard_paper_KR.docx
    paper/CardioGuard_paper_EN.docx
"""
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write(
        "[!] python-docx 未安装。请先运行: pip install python-docx\n"
    )
    raise

HERE = Path(__file__).parent
OUT_DIR = HERE / "paper"
OUT_DIR.mkdir(exist_ok=True)

LANG_FONTS = {
    "ZH": ("等线", "Calibri"),       # CJK fallback / EN
    "KR": ("Malgun Gothic", "Calibri"),
    "EN": ("Times New Roman", "Times New Roman"),
}

LANG_TITLES = {
    "ZH": "CardioGuard Industrial — 中文论文",
    "KR": "CardioGuard Industrial — 한국어 논문",
    "EN": "CardioGuard Industrial — English Paper",
}


def _set_run_font(run, cjk_font: str, ascii_font: str, size: float = 11):
    run.font.name = ascii_font
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), ascii_font)
    rFonts.set(qn("w:hAnsi"), ascii_font)
    rFonts.set(qn("w:eastAsia"), cjk_font)


def _add_para(doc: Document, text: str, cjk: str, asc: str,
              size=11, bold=False, italic=False, color=None,
              align=None, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    # split inline bold **...** + code `...` minimal handling
    segments = _inline_split(text)
    for seg_text, seg_bold, seg_italic, seg_code in segments:
        run = p.add_run(seg_text)
        run.bold = bold or seg_bold
        run.italic = italic or seg_italic
        _set_run_font(run, cjk, "Consolas" if seg_code else asc, size)
        if color:
            run.font.color.rgb = color
    return p


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def _inline_split(text: str):
    """Return [(text, bold, italic, code)] segments."""
    out = []
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            out.append((text[pos:m.start()], False, False, False))
        tok = m.group(0)
        if tok.startswith("**"):
            out.append((tok[2:-2], True, False, False))
        elif tok.startswith("*"):
            out.append((tok[1:-1], False, True, False))
        elif tok.startswith("`"):
            out.append((tok[1:-1], False, False, True))
        pos = m.end()
    if pos < len(text):
        out.append((text[pos:], False, False, False))
    return out


def _add_heading(doc, text: str, level: int, cjk: str, asc: str):
    sizes = {1: 22, 2: 16, 3: 13, 4: 12}
    colors = {1: RGBColor(0x0F, 0x17, 0x2A),
              2: RGBColor(0xDC, 0x26, 0x26),
              3: RGBColor(0x0F, 0x17, 0x2A),
              4: RGBColor(0x33, 0x41, 0x55)}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    _set_run_font(run, cjk, asc, sizes.get(level, 12))
    run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))
    p.space_after = Pt(6)
    return p


def _add_table_md(doc, header_cells, rows, cjk, asc):
    table = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(header_cells):
        run = hdr[i].paragraphs[0].add_run(h.strip())
        run.bold = True
        _set_run_font(run, cjk, asc, 10)
    for r_i, row in enumerate(rows, start=1):
        for c_i, cell in enumerate(row):
            if c_i >= len(header_cells):
                continue
            paragraph = table.rows[r_i].cells[c_i].paragraphs[0]
            for seg_text, b, i_, code in _inline_split(cell.strip()):
                run = paragraph.add_run(seg_text)
                run.bold = b
                run.italic = i_
                _set_run_font(run, cjk, "Consolas" if code else asc, 10)


def _is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$", line))


def md_to_docx(md_text: str, out_path: Path, lang: str):
    cjk_font, asc_font = LANG_FONTS[lang]
    doc = Document()

    # 全局默认样式
    style = doc.styles["Normal"]
    style.font.name = asc_font
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), asc_font)
    rFonts.set(qn("w:hAnsi"), asc_font)
    rFonts.set(qn("w:eastAsia"), cjk_font)

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        # Skip horizontal rules / front-matter separators
        if line in ("---", "***", "___"):
            doc.add_paragraph().add_run("").add_break()
            i += 1
            continue
        # Headings
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            _add_heading(doc, m.group(2).strip(), level, cjk_font, asc_font)
            i += 1
            continue
        # Code block
        if line.startswith("```"):
            j = i + 1
            block = []
            while j < len(lines) and not lines[j].startswith("```"):
                block.append(lines[j])
                j += 1
            for bl in block:
                p = doc.add_paragraph()
                run = p.add_run(bl)
                _set_run_font(run, cjk_font, "Consolas", 9.5)
                p.paragraph_format.left_indent = Cm(0.5)
            i = j + 1
            continue
        # Table
        if "|" in line and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
            header = [c for c in line.strip().strip("|").split("|")]
            rows = []
            j = i + 2
            while j < len(lines) and "|" in lines[j]:
                r = [c for c in lines[j].strip().strip("|").split("|")]
                rows.append(r)
                j += 1
            _add_table_md(doc, header, rows, cjk_font, asc_font)
            doc.add_paragraph()
            i = j
            continue
        # Bullet list
        if re.match(r"^\s*[-*+]\s+", line):
            text = re.sub(r"^\s*[-*+]\s+", "", line)
            _add_para(doc, text, cjk_font, asc_font,
                      style="List Bullet")
            i += 1
            continue
        # Ordered list
        if re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            _add_para(doc, text, cjk_font, asc_font,
                      style="List Number")
            i += 1
            continue
        # Blank line
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue
        # Default paragraph
        _add_para(doc, line, cjk_font, asc_font)
        i += 1

    doc.save(out_path)


def build():
    targets = [
        ("ZH", HERE / "PAPER_ZH.md", OUT_DIR / "CardioGuard_paper_ZH.docx"),
        ("KR", HERE / "PAPER_KR.md", OUT_DIR / "CardioGuard_paper_KR.docx"),
        ("EN", HERE / "PAPER_EN.md", OUT_DIR / "CardioGuard_paper_EN.docx"),
    ]
    sys.stdout.reconfigure(encoding="utf-8")
    for lang, src, dst in targets:
        if not src.exists():
            print(f"[!] {src.name} missing — skip {lang}")
            continue
        md = src.read_text(encoding="utf-8")
        md_to_docx(md, dst, lang)
        kb = dst.stat().st_size / 1024
        print(f"[OK] {lang}: {dst.relative_to(HERE)}  ({kb:.1f} KB)")


if __name__ == "__main__":
    build()
